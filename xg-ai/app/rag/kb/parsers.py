"""File → plain text. Pluggable per extension. PDF / DOCX deps are imported
lazily so the sidecar still starts when they're missing — caller gets a clear
error at upload time instead of an opaque import error at boot."""
from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_to_text(filename: str, content: bytes) -> str:
    """Returns the document's plain-text representation. Caller decides how
    to chunk afterwards. Raises ValueError on unsupported extension and
    RuntimeError when a parser dependency is missing."""
    suffix = Path(filename).suffix.lower().lstrip(".")
    if suffix in ("md", "markdown", "txt", "log"):
        return _decode_text(content)
    if suffix == "pdf":
        return _parse_pdf(content)
    if suffix in ("docx",):
        return _parse_docx(content)
    raise ValueError(f"unsupported file extension: .{suffix}")


def _decode_text(content: bytes) -> str:
    # try utf-8 first, then fall back to gbk for legacy Chinese files
    for enc in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _parse_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader  # noqa: WPS433
    except ImportError as e:
        raise RuntimeError(
            "PDF 解析需要 pypdf：pip install pypdf"
        ) from e
    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception as e:  # pragma: no cover — pdf parser is brittle
            logger.warning("pdf page parse failed: %s", e)
            text = ""
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _parse_docx(content: bytes) -> str:
    try:
        import docx  # python-docx  # noqa: WPS433
    except ImportError as e:
        raise RuntimeError(
            "DOCX 解析需要 python-docx：pip install python-docx"
        ) from e
    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs if p.text]
    # tables — flatten cell text per row
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text)
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)
